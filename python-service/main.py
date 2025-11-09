from fastapi import FastAPI, File, UploadFile, HTTPException, Form, Query
from fastapi.responses import JSONResponse, Response
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from PIL import Image
import io
import magic
import PyPDF2
import docx
import imagehash
from typing import Optional, List
import base64

from ai import (
    AISettings,
    ChatManager,
    EmbeddingManager,
    Summarizer,
    Tagger,
    extract_text,
    get_settings,
)

settings = get_settings()
summarizer = Summarizer(settings)
tagger = Tagger(settings)
embedding_manager = EmbeddingManager(settings)
chat_manager = ChatManager(settings)

app = FastAPI(title="LocalCloud Python Service", version="2.0.0")

# CORS middleware to allow Node.js backend to call this service
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5000", "http://localhost:3000"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


class ChatRequest(BaseModel):
    question: str
    top_k: int | None = 4


@app.get("/")
async def root():
    return {
        "service": "LocalCloud Python Service",
        "version": "2.0.0",
        "status": "running",
        "features": [
            "Image thumbnail generation",
            "File metadata extraction",
            "Text extraction from PDFs",
            "Text extraction from DOCX",
            "Duplicate image detection",
            "AI summarisation and tagging",
            "Semantic search",
            "Chat with your files",
        ],
    }


@app.get("/health")
async def health_check():
    return {"status": "healthy"}


@app.post("/ai/process")
async def process_file(
    file_id: str = Form(...),
    user_id: Optional[str] = Form(None),
    filename: Optional[str] = Form(None),
    file: UploadFile = File(...),
):
    try:
        contents = await file.read()
        text, mime_type = extract_text(contents, filename or file.filename)
        summary = await summarizer.summarize(text)
        tags = await tagger.generate_tags(text)
        embedding = await embedding_manager.upsert_document(
            file_id=file_id,
            text=text,
            metadata={
                "file_id": file_id,
                "user_id": user_id or "",
                "filename": filename or file.filename,
            },
        )

        return {
            "success": True,
            "file_id": file_id,
            "summary": summary,
            "tags": tags,
            "embedding": embedding,
            "text_preview": text[:1000],
            "mime_type": mime_type,
        }
    except Exception as exc:  # pragma: no cover - runtime failure
        raise HTTPException(status_code=400, detail=f"Failed to process file: {exc}")


@app.get("/ai/search")
async def semantic_search(q: str = Query(..., description="Search query"), top_k: int = 5):
    try:
        results = await embedding_manager.search(q, top_k=top_k)
        return {"success": True, "results": results}
    except Exception as exc:  # pragma: no cover
        raise HTTPException(status_code=400, detail=f"Semantic search failed: {exc}")


@app.delete("/ai/vector/{file_id}")
async def delete_vector(file_id: str):
    try:
        embedding_manager.delete_document(file_id)
        return {"success": True}
    except Exception as exc:  # pragma: no cover
        raise HTTPException(status_code=400, detail=f"Failed to delete vector: {exc}")


@app.post("/ai/chat")
async def chat_with_files(request: ChatRequest):
    try:
        response = await chat_manager.chat(request.question, top_k=request.top_k or 4)
        return {"success": True, **response}
    except Exception as exc:  # pragma: no cover
        raise HTTPException(status_code=400, detail=f"Chat failed: {exc}")


@app.post("/api/thumbnail")
async def generate_thumbnail(
    file: UploadFile = File(...), size: Optional[int] = 200
):
    """
    Generate a thumbnail for an uploaded image.
    Returns base64 encoded thumbnail.
    """
    try:
        contents = await file.read()
        image = Image.open(io.BytesIO(contents))

        # Convert RGBA to RGB if necessary
        if image.mode in ("RGBA", "LA", "P"):
            background = Image.new("RGB", image.size, (255, 255, 255))
            if image.mode == "P":
                image = image.convert("RGBA")
            background.paste(image, mask=image.split()[-1] if image.mode in ("RGBA", "LA") else None)
            image = background

        # Create thumbnail
        image.thumbnail((size, size), Image.Resampling.LANCZOS)

        # Save to bytes
        thumb_io = io.BytesIO()
        image.save(thumb_io, format="JPEG", quality=85)
        thumb_io.seek(0)

        # Encode to base64
        thumb_base64 = base64.b64encode(thumb_io.read()).decode("utf-8")

        return {
            "success": True,
            "thumbnail": f"data:image/jpeg;base64,{thumb_base64}",
            "original_size": {"width": image.width, "height": image.height},
        }

    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to generate thumbnail: {str(e)}")


@app.post("/api/metadata")
async def extract_metadata(file: UploadFile = File(...)):
    """
    Extract comprehensive metadata from uploaded file.
    """
    try:
        contents = await file.read()

        # Detect MIME type
        mime = magic.Magic(mime=True)
        mime_type = mime.from_buffer(contents)

        metadata = {
            "filename": file.filename,
            "size": len(contents),
            "mime_type": mime_type,
            "file_type": mime_type.split("/")[0] if "/" in mime_type else "unknown",
        }

        # Extract image-specific metadata
        if mime_type.startswith("image/"):
            try:
                image = Image.open(io.BytesIO(contents))
                metadata["image_info"] = {
                    "width": image.width,
                    "height": image.height,
                    "format": image.format,
                    "mode": image.mode,
                }

                # Extract EXIF data if available
                exif = image.getexif()
                if exif:
                    metadata["exif"] = {k: str(v) for k, v in exif.items() if k in [271, 272, 274, 306]}

            except Exception as e:
                metadata["image_error"] = str(e)

        return {"success": True, "metadata": metadata}

    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to extract metadata: {str(e)}")


@app.post("/api/extract-text/pdf")
async def extract_text_from_pdf(file: UploadFile = File(...)):
    """
    Extract text content from PDF files.
    """
    try:
        contents = await file.read()
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(contents))

        text_content = []
        for page_num, page in enumerate(pdf_reader.pages):
            text = page.extract_text()
            text_content.append({"page": page_num + 1, "text": text})

        return {
            "success": True,
            "filename": file.filename,
            "total_pages": len(pdf_reader.pages),
            "content": text_content,
            "preview": text_content[0]["text"][:500] if text_content else "",
        }

    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to extract text from PDF: {str(e)}")


@app.post("/api/extract-text/docx")
async def extract_text_from_docx(file: UploadFile = File(...)):
    """
    Extract text content from DOCX files.
    """
    try:
        contents = await file.read()
        doc = docx.Document(io.BytesIO(contents))

        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

        full_text = "\n".join(paragraphs)

        return {
            "success": True,
            "filename": file.filename,
            "total_paragraphs": len(paragraphs),
            "content": paragraphs,
            "preview": full_text[:500],
        }

    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to extract text from DOCX: {str(e)}")


@app.post("/api/image-hash")
async def calculate_image_hash(file: UploadFile = File(...)):
    """
    Calculate perceptual hash for duplicate image detection.
    Uses average hash, difference hash, and perceptual hash.
    """
    try:
        contents = await file.read()
        image = Image.open(io.BytesIO(contents))

        # Calculate multiple hashes for better accuracy
        avg_hash = str(imagehash.average_hash(image))
        diff_hash = str(imagehash.dhash(image))
        percept_hash = str(imagehash.phash(image))

        return {
            "success": True,
            "filename": file.filename,
            "hashes": {
                "average_hash": avg_hash,
                "difference_hash": diff_hash,
                "perceptual_hash": percept_hash,
            },
        }

    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to calculate image hash: {str(e)}")


@app.post("/api/analyze-file")
async def analyze_file_comprehensive(file: UploadFile = File(...)):
    """
    Comprehensive file analysis combining multiple features.
    Returns metadata, thumbnail (for images), and text content (for documents).
    """
    try:
        contents = await file.read()

        # Detect MIME type
        mime = magic.Magic(mime=True)
        mime_type = mime.from_buffer(contents)

        result = {
            "success": True,
            "filename": file.filename,
            "size": len(contents),
            "mime_type": mime_type,
            "file_type": mime_type.split("/")[0] if "/" in mime_type else "unknown",
        }

        # Handle images
        if mime_type.startswith("image/"):
            try:
                image = Image.open(io.BytesIO(contents))
                result["image_info"] = {
                    "width": image.width,
                    "height": image.height,
                    "format": image.format,
                    "mode": image.mode,
                }

                # Generate thumbnail
                if image.mode in ("RGBA", "LA", "P"):
                    background = Image.new("RGB", image.size, (255, 255, 255))
                    if image.mode == "P":
                        image = image.convert("RGBA")
                    background.paste(image, mask=image.split()[-1] if image.mode in ("RGBA", "LA") else None)
                    image = background

                image.thumbnail((200, 200), Image.Resampling.LANCZOS)
                thumb_io = io.BytesIO()
                image.save(thumb_io, format="JPEG", quality=85)
                thumb_io.seek(0)
                result["thumbnail"] = f"data:image/jpeg;base64,{base64.b64encode(thumb_io.read()).decode('utf-8')}"

            except Exception as e:
                result["image_error"] = str(e)

        # Handle PDFs
        elif mime_type == "application/pdf":
            try:
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(contents))
                first_page = pdf_reader.pages[0].extract_text() if len(pdf_reader.pages) > 0 else ""
                result["document_info"] = {
                    "type": "pdf",
                    "total_pages": len(pdf_reader.pages),
                    "preview": first_page[:500],
                }
            except Exception as e:
                result["document_error"] = str(e)

        # Handle DOCX
        elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            try:
                doc = docx.Document(io.BytesIO(contents))
                paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
                full_text = "\n".join(paragraphs)
                result["document_info"] = {
                    "type": "docx",
                    "total_paragraphs": len(paragraphs),
                    "preview": full_text[:500],
                }
            except Exception as e:
                result["document_error"] = str(e)

        return result

    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to analyze file: {str(e)}")


if __name__ == "__main__":
    import uvicorn

    uvicorn.run(app, host="0.0.0.0", port=8000)

